from typing import BinaryIO, Optional
from langchain_community.document_loaders import PyPDFLoader, TextLoader
from langchain.text_splitter import RecursiveCharacterTextSplitter
import tempfile
import os
from docx import Document

from src.utils.config import config
from src.utils.logger import log_error
from src.models.document import DocumentType


class DocumentProcessor:
    def __init__(self):
        self.text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=1000,
            chunk_overlap=200,
            length_function=len,
        )

    def extract_text_from_pdf(self, file_content: bytes, filename: str) -> str:
        """Extract text from PDF using LangChain PyPDFLoader."""
        try:
            # Save uploaded file to temporary location
            with tempfile.NamedTemporaryFile(delete=False, suffix=".pdf") as temp_file:
                temp_file.write(file_content)
                temp_file_path = temp_file.name

            # Use LangChain PyPDFLoader
            loader = PyPDFLoader(temp_file_path)
            documents = loader.load()

            # Combine all pages
            text = "\n".join([doc.page_content for doc in documents])

            # Clean up temporary file
            os.unlink(temp_file_path)

            return text

        except Exception as e:
            log_error(f"Error extracting text from PDF: {str(e)}")
            return ""

    def extract_text_from_txt(self, file_content: bytes, filename: str) -> str:
        """Extract text from TXT file."""
        try:
            # Try different encodings
            encodings = ["utf-8", "utf-16", "latin-1", "cp1252"]

            for encoding in encodings:
                try:
                    text = file_content.decode(encoding)
                    return text
                except UnicodeDecodeError:
                    continue

            # If all encodings fail, use utf-8 with error handling
            return file_content.decode("utf-8", errors="ignore")

        except Exception as e:
            log_error(f"Error extracting text from TXT: {str(e)}")
            return ""

    def extract_text_from_docx(self, file_content: bytes, filename: str) -> str:
        """Extract text from DOCX file."""
        try:
            # Save uploaded file to temporary location
            with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as temp_file:
                temp_file.write(file_content)
                temp_file_path = temp_file.name

            # Use python-docx to extract text
            from docx import Document as DocxDocument

            doc = DocxDocument(temp_file_path)

            # Extract text from all paragraphs
            text_parts = []
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_parts.append(paragraph.text)

            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            text_parts.append(cell.text)

            # Clean up temporary file
            os.unlink(temp_file_path)

            # Join all text parts
            full_text = "\n".join(text_parts)
            return full_text

        except Exception as e:
            log_error(f"Error extracting text from DOCX: {str(e)}")
            return ""

    def extract_text(self, file_content: bytes, filename: str) -> str:
        """Extract text based on file extension."""
        file_ext = filename.lower().split(".")[-1]

        if file_ext == "pdf":
            return self.extract_text_from_pdf(file_content, filename)
        elif file_ext == "txt":
            return self.extract_text_from_txt(file_content, filename)
        elif file_ext in ["docx", "doc"]:
            return self.extract_text_from_docx(file_content, filename)
        else:
            log_error(f"Unsupported file type: {file_ext}")
            return ""

    def split_text_into_chunks(self, text: str) -> list:
        """Split text into manageable chunks for processing."""
        return self.text_splitter.split_text(text)

    def detect_document_type(self, text: str) -> DocumentType:
        """Detect document type based on content."""
        text_lower = text.lower()

        # Rental agreement keywords
        rental_keywords = [
            "lease",
            "rent",
            "tenant",
            "landlord",
            "property",
            "premises",
            "deposit",
        ]

        # Loan agreement keywords
        loan_keywords = [
            "loan",
            "borrow",
            "lender",
            "principal",
            "interest",
            "repayment",
            "credit",
        ]

        # Employment keywords
        employment_keywords = [
            "employment",
            "employee",
            "employer",
            "salary",
            "wages",
            "position",
            "job",
        ]

        # NDA keywords
        nda_keywords = ["confidential", "non-disclosure", "proprietary", "trade secret"]

        # Service agreement keywords
        service_keywords = [
            "service",
            "provider",
            "client",
            "deliverables",
            "scope of work",
        ]

        # Count keyword matches
        scores = {
            DocumentType.RENTAL: sum(
                1 for keyword in rental_keywords if keyword in text_lower
            ),
            DocumentType.LOAN: sum(
                1 for keyword in loan_keywords if keyword in text_lower
            ),
            DocumentType.EMPLOYMENT: sum(
                1 for keyword in employment_keywords if keyword in text_lower
            ),
            DocumentType.NDA: sum(
                1 for keyword in nda_keywords if keyword in text_lower
            ),
            DocumentType.SERVICE: sum(
                1 for keyword in service_keywords if keyword in text_lower
            ),
        }

        # Return type with highest score, or OTHER if no clear match
        if max(scores.values()) > 2:
            return max(scores, key=scores.get)
        else:
            return DocumentType.OTHER

    def extract_metadata(self, text: str) -> dict:
        """Extract metadata from document text."""
        metadata = {
            "word_count": len(text.split()),
            "character_count": len(text),
            "estimated_reading_time": len(text.split()) // 200,  # Assuming 200 WPM
        }

        return metadata
